"""
Smart Import pipeline.

Handles: .xlsx/.xls, .csv, .docx (tables), .pdf (tables)
Produces a normalized list[dict] of rows, then validates/dedupes against
the DB before anything gets committed. Nothing is written until
commit_import() is called with the caller's approval.
"""
import re
import io
import datetime
import pandas as pd
import docx
import pdfplumber


# --- canonical field names & the header aliases we'll accept -------------

STUDENT_FIELD_ALIASES = {
    "register_number": ["register number", "reg no", "reg. no", "regno", "register_no", "usn"],
    "name": ["student name", "name", "full name"],
    "department": ["department", "dept"],
    "course": ["course"],
    "section": ["section", "sec"],
    "academic_year": ["academic year", "year", "batch"],
    "gender": ["gender", "sex"],
    "date_of_birth": ["date of birth", "dob"],
    "mobile_number": ["mobile number", "mobile", "phone", "contact number"],
    "email": ["email", "email id", "e-mail"],
    "address": ["address"],
    "cgpa": ["cgpa"],
    "percentage": ["percentage", "%"],
    "backlogs": ["backlogs", "arrears"],
    "skills": ["skills"],
    "resume_link": ["resume link", "resume", "resume url"],
    "placement_status": ["placement status", "status"],
    "eligible_status": ["eligible status", "eligible"],
}

COMPANY_FIELD_ALIASES = {
    "name": ["company name", "name", "company"],
    "industry": ["industry", "sector"],
    "state": ["state"],
    "location": ["location", "city"],
    "hr_name": ["hr name", "hr contact person", "contact person"],
    "hr_email": ["hr email", "hr email id"],
    "hr_contact_number": ["hr contact number", "hr phone", "hr contact"],
    "visit_date": ["visit date", "date of visit"],
    "package_amount": ["package", "ctc"],
    "min_package": ["minimum package", "min package"],
    "max_package": ["maximum package", "max package"],
    "avg_package": ["average package", "avg package"],
    "eligible_departments": ["eligible departments", "departments"],
    "min_cgpa": ["minimum cgpa", "min cgpa"],
    "allowed_backlogs": ["allowed backlogs", "backlogs allowed"],
    "hiring_count": ["hiring count", "openings", "positions"],
    "students_selected": ["students selected", "selected students", "register numbers selected"],
}


def _normalize_header(h):
    return re.sub(r"[^a-z0-9]", "", str(h).strip().lower())


def _build_lookup(aliases):
    lookup = {}
    for canonical, alts in aliases.items():
        for a in [canonical.replace("_", " ")] + alts:
            lookup[_normalize_header(a)] = canonical
    return lookup


def map_columns(columns, aliases):
    """Return {original_column_name: canonical_field_name or None}."""
    lookup = _build_lookup(aliases)
    mapping = {}
    for col in columns:
        mapping[col] = lookup.get(_normalize_header(col))
    return mapping


# --- file readers ----------------------------------------------------------

def read_tabular_file(file_storage, filename):
    """Return a pandas DataFrame from xlsx/xls/csv/docx/pdf."""
    ext = filename.rsplit(".", 1)[-1].lower()
    raw = file_storage.read()

    if ext in ("xlsx", "xls"):
        return pd.read_excel(io.BytesIO(raw))

    if ext == "csv":
        return pd.read_csv(io.BytesIO(raw))

    if ext == "docx":
        d = docx.Document(io.BytesIO(raw))
        if not d.tables:
            raise ValueError("No table found in the Word document")
        table = d.tables[0]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        header, *data_rows = rows
        return pd.DataFrame(data_rows, columns=header)

    if ext == "pdf":
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            all_rows = []
            header = None
            for page in pdf.pages:
                for table in page.extract_tables():
                    if not table:
                        continue
                    if header is None:
                        header, *rows = table
                    else:
                        rows = table[1:] if table[0] == header else table
                    all_rows.extend(rows)
            if header is None:
                raise ValueError("No table found in the PDF")
            return pd.DataFrame(all_rows, columns=header)

    raise ValueError(f"Unsupported file type: .{ext}")


# --- validation --------------------------------------------------------

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
MOBILE_RE = re.compile(r"^\+?\d{10,13}$")


def _clean_val(v):
    if pd.isna(v):
        return None
    if isinstance(v, str):
        v = v.strip()
        return v if v else None
    return v


def _parse_date(v):
    if v is None:
        return None
    try:
        return pd.to_datetime(v).date().isoformat()
    except Exception:
        return None


def normalize_student_row(row, mapping):
    out = {}
    for orig_col, canonical in mapping.items():
        if canonical:
            out[canonical] = _clean_val(row.get(orig_col))

    errors = []
    if not out.get("register_number"):
        errors.append("Missing register number")
    if not out.get("name"):
        errors.append("Missing student name")
    if out.get("email") and not EMAIL_RE.match(str(out["email"])):
        errors.append(f"Invalid email format: {out['email']}")
    if out.get("mobile_number") and not MOBILE_RE.match(str(out["mobile_number"]).replace(" ", "")):
        errors.append(f"Invalid mobile number: {out['mobile_number']}")
    if out.get("cgpa") is not None:
        try:
            cgpa = float(out["cgpa"])
            if not (0 <= cgpa <= 10):
                errors.append(f"CGPA out of range: {cgpa}")
            out["cgpa"] = cgpa
        except (ValueError, TypeError):
            errors.append(f"Invalid CGPA: {out['cgpa']}")
    if out.get("date_of_birth"):
        out["date_of_birth"] = _parse_date(out["date_of_birth"])
    if out.get("backlogs") is not None:
        try:
            out["backlogs"] = int(out["backlogs"])
        except (ValueError, TypeError):
            out["backlogs"] = 0

    return out, errors


def normalize_company_row(row, mapping):
    out = {}
    for orig_col, canonical in mapping.items():
        if canonical:
            out[canonical] = _clean_val(row.get(orig_col))

    errors = []
    if not out.get("name"):
        errors.append("Missing company name")
    if out.get("hr_email") and not EMAIL_RE.match(str(out["hr_email"])):
        errors.append(f"Invalid HR email format: {out['hr_email']}")
    if out.get("visit_date"):
        out["visit_date"] = _parse_date(out["visit_date"])
    for numeric_field in ("package_amount", "min_package", "max_package", "avg_package", "min_cgpa"):
        val = out.get(numeric_field)
        if val is not None:
            try:
                out[numeric_field] = float(val)
            except (ValueError, TypeError):
                errors.append(f"Invalid number for {numeric_field}: {val}")
                out[numeric_field] = None
    for int_field in ("allowed_backlogs", "hiring_count"):
        val = out.get(int_field)
        if val is not None:
            try:
                out[int_field] = int(float(val))
            except (ValueError, TypeError):
                out[int_field] = None

    return out, errors


def build_preview(df, kind):
    """
    kind: 'student' | 'company'
    Returns {mapping, rows: [{data, errors, is_duplicate}], summary}
    """
    aliases = STUDENT_FIELD_ALIASES if kind == "student" else COMPANY_FIELD_ALIASES
    normalizer = normalize_student_row if kind == "student" else normalize_company_row
    mapping = map_columns(list(df.columns), aliases)

    seen_keys = set()
    rows = []
    for _, raw_row in df.iterrows():
        data, errors = normalizer(raw_row, mapping)
        key = data.get("register_number") if kind == "student" else data.get("name")
        is_dup_in_file = key is not None and key in seen_keys
        if key is not None:
            seen_keys.add(key)
        rows.append({
            "data": data,
            "errors": errors,
            "duplicate_in_file": is_dup_in_file,
        })

    return {
        "column_mapping": mapping,
        "unmapped_columns": [c for c, v in mapping.items() if v is None],
        "rows": rows,
        "total_rows": len(rows),
        "rows_with_errors": sum(1 for r in rows if r["errors"]),
    }
